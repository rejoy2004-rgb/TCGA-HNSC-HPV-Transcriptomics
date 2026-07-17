import docx
import re
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = "Immune_Landscape_HNSC.docx"
try:
    doc = docx.Document(doc_path)
except Exception as e:
    print(f"Error opening document: {e}")
    sys.exit(1)

# 1. Parse generate_docx.py to extract the official references list
with open('generate_docx.py', 'r', encoding='utf-8') as f:
    code = f.read()

ref_match = re.search(r'references_list\s*=\s*\[(.*?)\]', code, re.DOTALL)
if not ref_match:
    print("Error: Could not find references_list in generate_docx.py")
    sys.exit(1)

ref_entries_str = ref_match.group(1)
ref_entries = re.findall(r'"(.*?)"', ref_entries_str)
if not ref_entries:
    ref_entries = re.findall(r"'(.*?)'", ref_entries_str)

# Map references to key author names and years for cross-referencing
ref_keys = []
for ref in ref_entries:
    # Match any word characters, spaces, or hyphens at the start of the reference, case-insensitively
    first_author_match = re.match(r'^([a-zA-Z\u00C0-\u00FF\s\-]+)', ref)
    if first_author_match:
        # Split by comma or space to get the first author's surname
        raw_author = first_author_match.group(1).strip()
        # Special cases
        if "zur Hausen" in ref:
            author = "zur Hausen"
        elif "Cancer Genome Atlas" in ref:
            author = "Cancer Genome Atlas Network"
        else:
            author = raw_author.split(',')[0].split(' ')[0].strip()
    else:
        author = "Unknown"
        
    year_match = re.search(r'\((\d{4})\)', ref)
    year = year_match.group(1) if year_match else "Unknown"
    ref_keys.append((author, year, ref))

print(f"Mapped {len(ref_keys)} references:")
for author, year, ref in ref_keys:
    print(f"  - {author} ({year})")

# 2. Extract plain text from all paragraphs in the document
# Exclude the "References" section itself to avoid self-matching
paragraphs = []
in_references = False
for p in doc.paragraphs:
    text = p.text.strip()
    if text == "References":
        in_references = True
    if in_references:
        continue
    if text:
        paragraphs.append(text)

# Also look inside table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text = p.text.strip()
                if text:
                    paragraphs.append(text)

full_text = " ".join(paragraphs)

# 3. Find in-text citations
citations_found = []

# Pattern 1: (Author, Year) or (Author et al., Year) or (Author & Author, Year)
# e.g., (Ang et al., 2010), (Love et al., 2014), (Wouters and Nelson, 2021)
pat1 = r'\(([^)]+?),\s*(\d{4})\)'
for m in re.finditer(pat1, full_text):
    authors, year = m.group(1), m.group(2)
    # Filter out common false positives
    if not authors.strip().startswith(("Figure", "Table", "Supplementary", "Stage", "PC", "log2", "FDR")):
        citations_found.append((authors.strip(), year.strip(), f"({authors}, {year})"))

# Pattern 2: Author et al. (Year) or Author (Year)
# e.g., Love et al. (2014), zur Hausen (2009), Newman et al. (2019)
pat2 = r'\b([a-zA-Z\u00C0-\u00FF\s\-\&\.]+?(?:\set\sal\.)?)\s*\((\d{4})\)'
for m in re.finditer(pat2, full_text):
    authors, year = m.group(1), m.group(2)
    if not authors.strip().startswith(("Figure", "Table", "Supplementary", "Stage", "PC")):
        citations_found.append((authors.strip(), year.strip(), f"{authors} ({year})"))

print(f"\nExtracted {len(citations_found)} potential in-text citations:")
unique_citations = set()
for auth, yr, raw in citations_found:
    unique_citations.add((auth, yr, raw))
    print(f"  - {raw}")

# Helper function to clean author strings for matching
def clean_author(name):
    # Lowercase
    n = name.lower()
    # Replace "et al." with empty string
    n = n.replace("et al.", "")
    # Replace "and" or "&" with space ONLY when they are separate words
    n = re.sub(r'\b(and|\&)\b', ' ', n)
    # Remove extra spaces
    n = re.sub(r'\s+', ' ', n).strip()
    return n

# 4. Cross-reference
print("\nCross-Referencing Citations:")
mismatches_found = False
matched_ref_indices = set()

for auth, yr, raw in sorted(unique_citations, key=lambda x: x[0]):
    matched = False
    for idx, (ref_auth, ref_yr, ref) in enumerate(ref_keys):
        if yr == ref_yr:
            # Check for author name matching
            c_auth = clean_author(auth)
            c_ref_auth = clean_author(ref_auth)
            
            auth_parts = re.split(r'[\s,]+', c_auth)
            ref_parts = re.split(r'[\s,]+', c_ref_auth)
            
            # Match if there's name similarity
            if any(p in ref_parts for p in auth_parts) or c_ref_auth in c_auth or c_auth in c_ref_auth:
                matched = True
                matched_ref_indices.add(idx)
                break
                
    if not matched:
        print(f"  [MISSING REFERENCE] Citation '{raw}' has no matching entry in reference list!")
        mismatches_found = True
    else:
        print(f"  [OK] Citation '{raw}' matches reference key '{ref_keys[list(matched_ref_indices)[-1]][0]} ({yr})'")

# Check for unused references
print("\nChecking for Unused References:")
for idx, (ref_auth, ref_yr, ref) in enumerate(ref_keys):
    if idx not in matched_ref_indices:
        print(f"  [UNUSED REFERENCE] '{ref_auth} ({ref_yr})' is in the reference list but not cited in text.")
        mismatches_found = True

if not mismatches_found:
    print("\nSUCCESS: All in-text citations are in the reference list, and all reference list entries are used in the text!")
else:
    print("\nWARNING: Some citation mismatches were found. Please verify.")
