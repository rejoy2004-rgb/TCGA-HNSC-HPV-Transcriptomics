import docx
import os

doc_paths = ["Immune_Landscape_HNSC.docx"]
if os.path.exists("Immune_Landscape_HNSC_Updated.docx"):
    doc_paths.append("Immune_Landscape_HNSC_Updated.docx")

for doc_path in doc_paths:
    if not os.path.exists(doc_path):
        continue
    
    print(f"\n=========================================")
    print(f"Verifying Document Content ({doc_path}):")
    print(f"=========================================")
    
    try:
        doc = docx.Document(doc_path)
        
        # Check for fallback text placeholders
        placeholders = []
        for p in doc.paragraphs:
            if "[Placeholder for Figure:" in p.text:
                placeholders.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "[Placeholder for Figure:" in p.text:
                            placeholders.append(p.text.strip())

        if placeholders:
            print(f"  WARNING: Found {len(placeholders)} fallback placeholders in text:")
            for pl in placeholders:
                print(f"    - {pl}")
        else:
            print("  Success: No fallback placeholders found in the document!")

        # Count inline shapes / pictures
        shapes_count = 0
        for p in doc.paragraphs:
            if "pic:pic" in p._p.xml:
                shapes_count += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "pic:pic" in p._p.xml:
                            shapes_count += 1

        print(f"  Total Embedded Images Found: {shapes_count}")

        # Check that the file exists and is of correct size
        size_bytes = os.path.getsize(doc_path)
        print(f"  Document File Size: {size_bytes} bytes ({size_bytes / 1024 / 1024:.2f} MB)")
    except Exception as e:
        print(f"  ERROR verifying document: {e}")
