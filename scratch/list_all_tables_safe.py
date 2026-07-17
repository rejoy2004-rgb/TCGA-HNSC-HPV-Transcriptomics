import docx

doc_path = "Immune_Landscape_HNSC.docx"
doc = docx.Document(doc_path)

print("Listing all Tables in HNSC docx:")
table_count = 0
for idx, table in enumerate(doc.tables):
    table_count += 1
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    
    # Safely encode text to print on cp1252 terminal
    safe_row = [c.encode('ascii', errors='replace').decode('ascii') for c in first_row]
    
    caption = "No caption found"
    for p in doc.paragraphs:
        if f"Table {table_count}" in p.text:
            caption = p.text.strip().encode('ascii', errors='replace').decode('ascii')
            break
            
    print(f"Table {table_count}:")
    print(f"  Caption: {caption[:120]}...")
    print(f"  Header: {safe_row}")
