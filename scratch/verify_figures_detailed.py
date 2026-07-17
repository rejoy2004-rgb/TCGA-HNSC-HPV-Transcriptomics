import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = "Immune_Landscape_HNSC.docx"
if not os.path.exists(doc_path):
    print(f"Error: {doc_path} does not exist.")
    sys.exit(1)

doc = docx.Document(doc_path)
print(f"Detailed Figure Check for {doc_path}:\n")

# Find all paragraphs with images and their text
fig_count = 0
supp_count = 0
paragraphs_with_images = []

# We can search the XML of each paragraph to check if it contains a picture
for i, p in enumerate(doc.paragraphs):
    xml_str = p._p.xml
    if "pic:pic" in xml_str:
        # This paragraph contains an image!
        # Let's find the nearest caption text (usually the next paragraph or the paragraph itself)
        caption = ""
        # Look at the next few paragraphs
        for next_p in doc.paragraphs[i:i+3]:
            text = next_p.text.strip()
            if text.startswith(("Figure ", "Supplementary Figure ")):
                caption = text
                break
        
        paragraphs_with_images.append((i, caption))
        if "Supplementary Figure" in caption:
            supp_count += 1
        elif "Figure" in caption:
            fig_count += 1

print(f"Summary of Images Found:")
print(f" - Main Figures (Figure 1 to 16): {fig_count}")
print(f" - Supplementary Figures (S1 to S3): {supp_count}")
print(f" - Total Images Embedded: {len(paragraphs_with_images)}")

print("\nList of Embedded Figures and Captions:")
for idx, (p_idx, caption) in enumerate(paragraphs_with_images, 1):
    if caption:
        print(f"  {idx}. [Paragraph {p_idx}] Caption: {caption}")
    else:
        # Check if we can find any text in the paragraph itself or preceding it
        nearby_text = ""
        for offset in range(-2, 3):
            if 0 <= p_idx + offset < len(doc.paragraphs):
                t = doc.paragraphs[p_idx + offset].text.strip()
                if t:
                    nearby_text += f" [p{p_idx+offset}]: {t[:80]}..."
        print(f"  {idx}. [Paragraph {p_idx}] Image without adjacent Figure caption. Nearby text: {nearby_text}")

# Also check inside tables just in case images are placed inside cells
table_images = 0
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, p in enumerate(cell.paragraphs):
                if "pic:pic" in p._p.xml:
                    table_images += 1
                    print(f"  * Found image in Table {t_idx+1}, Row {r_idx+1}, Cell {c_idx+1}")

if table_images:
    print(f" - Total Table Images: {table_images}")
