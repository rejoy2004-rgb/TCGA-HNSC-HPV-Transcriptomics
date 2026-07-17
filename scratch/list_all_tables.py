import docx

doc_path = "Immune_Landscape_HNSC.docx"
doc = docx.Document(doc_path)

print("Listing all Tables in HNSC docx:")
table_count = 0
for idx, table in enumerate(doc.tables):
    # Print the first row of each table to identify it
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    table_count += 1
    # Try to find a paragraph preceding the table (often has the Table caption)
    caption = "No caption found"
    # Find paragraph index of table in the body
    for p_idx, p in enumerate(doc.paragraphs):
        # python-docx doesn't easily map paragraphs to preceding tables directly, 
        # but we can look for paragraphs containing "Table X"
        if f"Table {table_count}" in p.text:
            caption = p.text.strip()
            break
            
    print(f"Table {table_count}:")
    print(f"  Caption: {caption}")
    print(f"  Header: {first_row}")
