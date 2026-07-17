import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Sets background color of a cell (HEX format, e.g., 'F2F2F2')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a cell in twentieths of a point (dxa)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def set_cell_borders(cell, **kwargs):
    """
    Sets individual cell borders.
    kwargs can be: top, bottom, left, right, insideH, insideV
    Value should be a dict: {'sz': 4, 'val': 'single', 'color': 'D3D3D3', 'space': '0'}
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), edge_data.get('val', 'single'))
            border.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            border.set(qn('w:space'), str(edge_data.get('space', 0)))
            border.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(border)

def add_page_number(run):
    """Inserts a page number field into a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_total_page_number(run):
    """Inserts a NUMPAGES field into a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_figure(doc, image_name, caption_text, width_inches=4.5):
    """Adds centered image and caption styled professionally to the document. Reserves a space gap if not found."""
    if os.path.exists(image_name):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(4)
        p_img.paragraph_format.keep_with_next = True
        r_img = p_img.add_run()
        r_img.add_picture(image_name, width=Inches(width_inches))
    else:
        # Create a professional placeholder table representing a space gap for the figure
        border_format = {'sz': 4, 'val': 'single', 'color': 'D3D3D3', 'space': '0'}
        placeholder_table = doc.add_table(rows=1, cols=1)
        placeholder_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = placeholder_table.cell(0, 0)
        set_cell_borders(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
        set_cell_margins(cell, top=500, bottom=500, left=500, right=500)
        set_cell_background(cell, "FAFAFA")
        
        # Set cell width and height
        cell.width = Inches(width_inches)
        placeholder_table.rows[0].height = Inches(3.0)
        
        p_fallback = cell.paragraphs[0]
        p_fallback.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fallback.paragraph_format.space_before = Pt(50)  # Spacing to center text vertically in the 3-inch box
        p_fallback.paragraph_format.space_after = Pt(0)
        r_fallback = p_fallback.add_run(f"[Placeholder for Figure: {os.path.basename(image_name)}]\n(A space gap of 3.0 inches is reserved here for figure attachment)")
        r_fallback.font.italic = True
        r_fallback.font.bold = True
        r_fallback.font.size = Pt(10)
        r_fallback.font.color.rgb = RGBColor(128, 128, 128)
        r_fallback.font.name = 'Times New Roman'
        
        # Add spacing paragraph after the table
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(4)
        p_space.paragraph_format.space_after = Pt(0)
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(16)
    p_cap.paragraph_format.keep_with_next = True
    
    parts = caption_text.split(".", 1)
    if len(parts) == 2:
        r_lbl = p_cap.add_run(parts[0] + ".")
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.color.rgb = RGBColor(80, 80, 80)
        r_lbl.font.name = 'Times New Roman'
        
        r_desc = p_cap.add_run(parts[1])
        r_desc.font.size = Pt(9.5)
        r_desc.font.color.rgb = RGBColor(80, 80, 80)
        r_desc.font.name = 'Times New Roman'
    else:
        r_cap = p_cap.add_run(caption_text)
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = RGBColor(80, 80, 80)
        r_cap.font.name = 'Times New Roman'

def main():
    doc = Document()
    
    # 1. Page margins (1 inch / 2.54 cm all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Configure Header & Footer
        # Header text
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_p.add_run("Immune and Transcriptional Differences in HNSC")
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(9)
        header_run.font.italic = True
        header_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer text
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run_text = footer_p.add_run("Page ")
        footer_run_text.font.name = 'Times New Roman'
        footer_run_text.font.size = Pt(10)
        footer_run_text.font.color.rgb = RGBColor(128, 128, 128)
        
        footer_num_run = footer_p.add_run()
        footer_num_run.font.name = 'Times New Roman'
        footer_num_run.font.size = Pt(10)
        footer_num_run.font.color.rgb = RGBColor(128, 128, 128)
        add_page_number(footer_num_run)
        
        footer_of_run = footer_p.add_run(" of ")
        footer_of_run.font.name = 'Times New Roman'
        footer_of_run.font.size = Pt(10)
        footer_of_run.font.color.rgb = RGBColor(128, 128, 128)
        
        footer_total_run = footer_p.add_run()
        footer_total_run.font.name = 'Times New Roman'
        footer_total_run.font.size = Pt(10)
        footer_total_run.font.color.rgb = RGBColor(128, 128, 128)
        add_total_page_number(footer_total_run)

    # 2. Configure Typography Styles
    # Normal style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(8)
    style_normal.paragraph_format.space_before = Pt(0)
    
    # Heading 1 style
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(14)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(30, 30, 30)
    style_h1.paragraph_format.space_before = Pt(16)
    style_h1.paragraph_format.space_after = Pt(6)
    style_h1.paragraph_format.keep_with_next = True
    style_h1.paragraph_format.line_spacing = 1.15
    
    # Heading 2 style
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(12)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(40, 40, 40)
    style_h2.paragraph_format.space_before = Pt(12)
    style_h2.paragraph_format.space_after = Pt(4)
    style_h2.paragraph_format.keep_with_next = True
    style_h2.paragraph_format.line_spacing = 1.15

    # TITLE PAGE
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(
        "Immune and Transcriptional Differences Between HPV-Positive and HPV-Negative Head and Neck Squamous Cell Carcinoma"
    )
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.name = 'Times New Roman'
    
    # Running Title
    p_rt = doc.add_paragraph()
    p_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rt.paragraph_format.space_after = Pt(36)
    run_rt_label = p_rt.add_run("Running Title: ")
    run_rt_label.font.bold = True
    run_rt_label.font.size = Pt(11)
    run_rt_val = p_rt.add_run("Immune and Transcriptional Differences in HNSC")
    run_rt_val.font.italic = True
    run_rt_val.font.size = Pt(11)
    
    # Abstract heading
    p_abs_h = doc.add_paragraph()
    p_abs_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_abs_h = p_abs_h.add_run("Abstract")
    run_abs_h.font.bold = True
    run_abs_h.font.size = Pt(13)
    p_abs_h.paragraph_format.space_after = Pt(6)
    
    # Structured Abstract body paragraph
    abstract_sections = [
        ("Background", "Human papillomavirus (HPV)-positive head and neck squamous cell carcinoma (HNSC) exhibits distinct clinical behaviour compared with HPV-negative disease, including superior survival outcomes and heightened response to immunotherapy. Characterizing the molecular and immunological differences between these cohorts is critical for developing personalized treatment strategies."),
        ("Objective", "This study tests whether HPV-positive HNSC tumors show a distinct immune and transcriptional profile compared with HPV-negative tumors."),
        ("Methods", "We performed an integrative transcriptomic analysis of 279 TCGA-HNSC samples (243 HPV-negative, 36 HPV-positive). We utilized DESeq2 on raw counts for differential expression, CIBERSORTx (relative mode, LM22 matrix) on normalized counts for immune deconvolution, pathway enrichment (GO/KEGG/GSEA), and Cox proportional-hazards survival modeling on variance-stabilized counts."),
        ("Results", "Differential expression identified 6,288 genes (FDR < 0.05, |log2FC| > 1), including meiotic/cancer-testis antigens (STAG3, SMC1B, ZFR2, RAD9B) upregulated in HPV-positive tumors, and keratinization markers (S100A7, DSC1, SPRR4) downregulated. CIBERSORTx revealed significant enrichment of plasma cells (median fraction 0.44 vs. 0.14, FDR = 0.001) and CD8+ T-cell bulk estimates (0.07 vs. 0.05, FDR = 0.020) in HPV-positive tumors, validated by CD8A/B correlation. Reciprocally, M0 macrophages were depleted, and M2 macrophages showed a non-significant downward trend. The log2-transformed CD8/M2 ratio was significantly higher in HPV-positive tumors (P = 7.43e-05). Grouped literature review linked upregulated immunoglobulin genes to a humoral/plasma-cell signature. Survival analysis identified ZFR2 as an exploratory prognostic marker associated with better survival (unstandardized HR = 0.9988, P < 0.001). Comparison with cervical cancer (CESC) revealed HNSC-specific survival patterns."),
        ("Conclusions", "HPV-positive HNSC exhibits a highly active immune microenvironment driven by B-cell, plasma-cell, and CD8+ T-cell infiltration, alongside suppressed keratinization programs. These transcriptomic distinctions reinforce the biological rationale for immunotherapy and highlight novel candidate biomarkers requiring prospective validation.")
    ]
    
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.paragraph_format.space_after = Pt(12)
    for title, text in abstract_sections:
        r_title = p_abs.add_run(f"{title}: ")
        r_title.bold = True
        r_text = p_abs.add_run(f"{text} ")
    
    # Keywords
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.line_spacing = 1.15
    r_kw_lbl = p_kw.add_run("Keywords: ")
    r_kw_lbl.bold = True
    r_kw_val = p_kw.add_run("Head and neck squamous cell carcinoma; Human papillomavirus; Tumour microenvironment; Immune deconvolution; Differential expression; CIBERSORTx; DESeq2; Plasma cells; CD8+ T cells; Immunoglobulin")
    p_kw.paragraph_format.space_after = Pt(12)
    
    p_abb = doc.add_paragraph()
    p_abb.paragraph_format.line_spacing = 1.15
    r_abb_lbl = p_abb.add_run("Abbreviations: ")
    r_abb_lbl.bold = True
    abbrev_list = [
        ("HPV", "Human papillomavirus"),
        ("HNSC", "Head and neck squamous cell carcinoma"),
        ("DEG", "Differentially expressed gene"),
        ("GO", "Gene Ontology"),
        ("KEGG", "Kyoto Encyclopedia of Genes and Genomes"),
        ("GSEA", "Gene Set Enrichment Analysis"),
        ("TME", "Tumour microenvironment"),
        ("TLS", "Tertiary lymphoid structures"),
        ("TCGA", "The Cancer Genome Atlas"),
        ("LM22", "Leukocyte signature matrix 22"),
        ("FDR", "False discovery rate"),
        ("NES", "Normalized enrichment score")
    ]
    abbrev_runs = [f"{abbrev}, {full}" for abbrev, full in abbrev_list]
    p_abb.add_run("; ".join(abbrev_runs) + ".")
    p_abb.paragraph_format.space_after = Pt(12)
    
    # Page Break after title/abstract
    doc.add_page_break()
    
    # 1. INTRODUCTION
    doc.add_heading("1. Introduction", level=1)
    
    doc.add_paragraph(
        "Head and neck squamous cell carcinoma (HNSC) represents the sixth most prevalent malignancy worldwide, "
        "accounting for approximately 900,000 new cases and 450,000 deaths annually (Johnson et al., 2020). "
        "The disease arises from the mucosal epithelia of the oral cavity, oropharynx, hypopharynx, and larynx, "
        "and is historically associated with tobacco and alcohol exposure. Over the past three decades, however, "
        "infection with high-risk human papillomavirus (HPV)—principally HPV-16—has emerged as an aetiologically "
        "distinct and increasingly dominant driver of oropharyngeal squamous cell carcinoma (Marur et al., 2016). "
        "HPV-positive HNSC now constitutes the majority of oropharyngeal cancers in high-income countries and "
        "differs fundamentally from its HPV-negative counterpart in epidemiology, molecular pathogenesis, and clinical "
        "behaviour (Ang et al., 2010)."
    )
    
    doc.add_paragraph(
        "The improved prognosis of HPV-positive HNSC—reflected in markedly superior overall and progression-free "
        "survival—has been linked to multiple factors, including younger patient age, lower mutational burden, "
        "and, critically, a more immune-permissive tumour microenvironment (TME) (Fakhry et al., 2008). "
        "The TME in HNSC is now recognised as a key determinant of therapeutic response and disease outcome, "
        "influencing sensitivity to radiation, chemotherapy, and immune checkpoint inhibitors (Ferris et al., 2016). "
        "Immune deconvolution studies using algorithms such as CIBERSORTx have demonstrated that HNSC harbours "
        "diverse immune infiltrates, yet the composition of these infiltrates differs substantially between HPV-positive "
        "and HPV-negative disease (Mandal et al., 2016). In particular, the roles of CD8+ T cells, B cells/plasma cells, "
        "and myeloid populations like macrophages are highly critical. CD8+ cytotoxic T cells are the primary mediators of "
        "cellular anti-tumor immunity, while tumor-infiltrating B cells and plasma cells have been shown to form tertiary "
        "lymphoid structures (TLS), promoting local humoral responses and predicting favorable outcomes (Helmink et al., 2020). "
        "Conversely, macrophages, especially of the undifferentiated M0 and M2 phenotypes, often promote an immunosuppressive TME "
        "that facilitates tumor evasion and metastasis."
    )
    
    doc.add_paragraph(
        "Despite accumulating clinical evidence, the transcriptomic and immunological landscape distinguishing "
        "HPV-associated from HPV-independent HNSC remains incompletely characterised at the population level. "
        "The Cancer Genome Atlas (TCGA) HNSC dataset offers an unprecedented opportunity to interrogate these "
        "differences systematically, with publicly available RNA-sequencing data and curated clinical annotations "
        "enabling high-powered, unbiased analyses (Cancer Genome Atlas Network, 2015)."
    )
    
    doc.add_paragraph(
        "This study tests whether HPV-positive HNSC tumors show a distinct immune and transcriptional profile compared with "
        "HPV-negative tumors. To address this objective, we integrated several analytical steps: acquiring TCGA expression data, "
        "stratifying samples by HPV status, performing CIBERSORTx immune cell deconvolution, conducting differential gene "
        "expression analysis (DEG), carrying out a structured literature review of top candidate genes, and performing Cox "
        "proportional-hazards survival modeling. We proposed the hypothesis that HPV-positive tumours would express "
        "a significantly immune-activated TME, represented by an enrichment of cytotoxic and humoral immune effector populations, "
        "opposite to a reciprocal suppression of keratinisation-associated transcriptional programmes characteristic of more aggressive "
        "squamous differentiation."
    )
    
    # 2. MATERIALS AND METHODS
    doc.add_heading("2. Materials and Methods", level=1)
    
    doc.add_heading("2.1 Dataset Acquisition and Sample Stratification", level=2)
    doc.add_paragraph(
        "Raw RNA-sequencing count data for the TCGA-HNSC cohort were obtained from the Genomic Data Commons (GDC) portal. "
        "HPV status annotations were derived from a validated published dataset that integrates HPV detection from RNA-seq "
        "reads with clinical metadata (Cancer Genome Atlas Network, 2015). To integrate the genomic, deconvolution, and clinical datasets, "
        "sample IDs were matched using the first 15 characters of the TCGA barcode. After exclusion of samples with ambiguous HPV status "
        "or missing survival data, the final analytical cohort comprised 279 samples: 243 HPV-negative and 36 HPV-positive cases "
        "(Table 1). Continuous variables were compared using the Wilcoxon rank-sum test, whereas categorical variables were compared "
        "using Fisher's exact test or Chi-square test as appropriate."
    )
    
    doc.add_heading("2.2 Gene Expression Data Format", level=2)
    doc.add_paragraph(
        "For differential gene expression (DEG) analysis, raw RNA-sequencing read counts were used directly as input for DESeq2, "
        "which incorporates internal normalization based on median-of-ratios. For immune deconvolution with CIBERSORTx, gene expression "
        "values were formatted as Transcripts Per Million (TPM) to provide standardized relative abundance across genes. For survival "
        "analyses, PCA visualization, and heatmap generation, gene expression counts were log-transformed after applying DESeq2's "
        "variance-stabilizing transformation (VST) to stabilize the variance across the range of mean values. These format specifications "
        "ensure comparability and reproducibility of the downstream analyses."
    )
    
    doc.add_heading("2.3 Differential Expression Analysis", level=2)
    doc.add_paragraph(
        "Differential expression between HPV-positive and HPV-negative tumours was performed using DESeq2 (v1.38; Love et al., 2014) "
        "in R (v4.3). Genes with low counts were removed by retaining only those with a row sum of at least 10 counts across all "
        "samples. Variance-stabilising transformation was applied for exploratory visualisation, including principal component "
        "analysis (PCA). Differential expression was assessed using the default DESeq2 negative binomial model with Wald testing, "
        "with HPV status as the primary covariate. Statistical significance thresholds were set at a Benjamini-Hochberg false "
        "discovery rate (FDR) < 0.05 and an absolute log2 fold change (|log2FC|) > 1. Results were visualised using a volcano plot "
        "with labelled top-ranked genes."
    )
    
    doc.add_heading("2.4 CIBERSORTx Immune Deconvolution", level=2)
    doc.add_paragraph(
        "Immune cell composition was estimated using CIBERSORTx with the LM22 signature matrix, which quantifies 22 immune cell "
        "populations from bulk RNA-sequencing data (Newman et al., 2019). Deconvolution was performed in relative mode without "
        "quantile normalisation, as recommended for RNA-seq data. No CIBERSORTx p-value filtering was applied to the samples, "
        "ensuring that all 279 patients were retained for an unbiased cohort comparison. Differences in cell-type fractions "
        "between HPV-positive and HPV-negative groups were assessed using the Wilcoxon rank-sum test, with Benjamini-Hochberg "
        "correction applied across all 22 cell types to account for multiple testing. A significance threshold of FDR < 0.05 was applied."
    )
    doc.add_paragraph(
        "To validate the fidelity of CIBERSORTx CD8+ T-cell estimates, Spearman rank correlations were computed between "
        "CIBERSORT-derived CD8 fractions and the expression levels of canonical CD8 T-cell markers CD8A and CD8B. "
        "Additionally, a CD8/M2 macrophage immune activation ratio was computed as log2((CD8 fraction + 1 × 10−4) / "
        "(M2 fraction + 1 × 10−4)) to provide a composite measure of cytotoxic versus immunosuppressive immune balance, using "
        "a pseudocount of 1e-4 to prevent numerical instability from zero values."
    )
    
    doc.add_heading("2.5 Pathway and Gene Set Enrichment Analysis", level=2)
    doc.add_paragraph(
        "Gene Ontology (GO) biological process and KEGG pathway enrichment analyses were conducted separately on upregulated "
        "and downregulated differentially expressed gene lists using the clusterProfiler package (Yu et al., 2012) in R. "
        "Background gene universe was defined as all genes passing the low-count filter. Statistical significance was "
        "defined as an adjusted P-value < 0.05."
    )
    doc.add_paragraph(
        "Gene Set Enrichment Analysis (GSEA) was performed on the complete pre-ranked gene list, ordered by the DESeq2 "
        "Wald statistic, using the fgsea package against the MSigDB Hallmark gene set collection. Significant enrichment "
        "was defined as an FDR-adjusted P-value < 0.25, in accordance with established GSEA reporting standards "
        "(Subramanian et al., 2005)."
    )
    
    doc.add_heading("2.6 Survival Analysis", level=2)
    doc.add_paragraph(
        "Univariate Cox proportional-hazards regression was performed for each gene using the survival package in R, with "
        "overall survival as the outcome. To address the scaling of hazard ratios, models were evaluated both on unstandardized "
        "expression values (which yield hazard ratios per single raw count change) and explained in the context of standardized "
        "expression (hazard ratio per 1 standard deviation increase in expression). Standardizing expression is essential for "
        "biological interpretation, as unstandardized counts lead to hazard ratios extremely close to unity due to the large range of "
        "expression counts. Patients were dichotomised into high and low expression groups at the median expression level for survival "
        "visualization. Kaplan-Meier survival curves were constructed for ZFR2, and the log-rank test was used to compare survival "
        "curves. These survival models are exploratory univariate analyses, as they do not adjust for clinical covariates "
        "(such as stage, age, and sex)."
    )
    
    # 3. RESULTS
    doc.add_heading("3. Results", level=1)
    
    doc.add_heading("3.1 Study Workflow and Patient Demographics", level=2)
    doc.add_paragraph(
        "The overall workflow of this bioinformatics analysis pipeline is schematically represented in Figure 1. "
        "The final analytical cohort comprised 279 TCGA-HNSC samples, including 243 HPV-negative and 36 HPV-positive tumours. "
        "The baseline clinical and demographic characteristics of the patients are summarized in Table 1."
    )
    
    # Insert Figure 1 (Workflow)
    add_figure(
        doc,
        "Figure_1_Workflow.png",
        "Figure 1. Schematic workflow of the bioinformatics analysis pipeline. TCGA-HNSC RNA-seq raw count data were stratified by HPV status and subjected to differential expression analysis (DESeq2), immune cell deconvolution (CIBERSORTx/LM22), pathway enrichment (GO/KEGG), GSEA, and survival analysis."
    )
    
    doc.add_paragraph(
        "Principal component analysis of variance-stabilised expression data demonstrated partial but incomplete separation "
        "of HPV-positive (teal) and HPV-negative (salmon) samples along PC1 (22% variance explained) and PC2 (14% variance "
        "explained), consistent with HPV status contributing substantially—though not exclusively—to global transcriptomic "
        "variability (Figure 2). The broad distribution along PC2 within both groups likely reflects the heterogeneity in "
        "tumour site, differentiation grade, and stromal content characteristic of HNSC."
    )
    
    # Insert Table 1 (Clinical Demographics) here
    p_t1_title = doc.add_paragraph()
    p_t1_title.paragraph_format.space_before = Pt(12)
    p_t1_title.paragraph_format.space_after = Pt(4)
    p_t1_title.paragraph_format.keep_with_next = True
    r_t1_title = p_t1_title.add_run("Table 1. Baseline Clinical and Demographic Characteristics of HNSC Patients Stratified by HPV Status")
    r_t1_title.font.bold = True
    r_t1_title.font.size = Pt(11)
    
    table_1_data = [
        ["Variable", "HPV+ (n = 36)", "HPV− (n = 243)", "P-value"],
        ["Age (years)", "", "", "0.053"],
        ["  Mean ± SD", "57.9 ± 10.8", "61.8 ± 12.4", ""],
        ["  Median (range)", "58.5 (35–82)", "62.0 (19–90)", ""],
        ["Gender, n (%)", "", "", "0.033"],
        ["  Female", "4 (11.1%)", "72 (29.6%)", ""],
        ["  Male", "32 (88.9%)", "171 (70.4%)", ""],
        ["Tumor Stage, n (%)", "", "", "0.845"],
        ["  Stage I", "2 (5.6%)", "12 (4.9%)", ""],
        ["  Stage II", "5 (13.9%)", "39 (16.0%)", ""],
        ["  Stage III", "4 (11.1%)", "34 (14.0%)", ""],
        ["  Stage IV", "12 (33.3%)", "132 (54.3%)", ""],
        ["  Unknown", "13 (36.1%)", "26 (10.7%)", ""],
        ["Smoking Status, n (%)", "", "", "0.234"],
        ["  Never Smoker", "10 (27.8%)", "42 (17.3%)", ""],
        ["  Ever Smoker", "26 (72.2%)", "194 (79.8%)", ""],
        ["  Unknown", "0 (0.0%)", "7 (2.9%)", ""],
        ["Primary Site, n (%)", "", "", "< 0.001"],
        ["  Oropharynx", "22 (61.1%)", "11 (4.5%)", ""],
        ["  Oral Cavity", "12 (33.3%)", "160 (65.8%)", ""],
        ["  Larynx", "1 (2.8%)", "71 (29.2%)", ""],
        ["  Hypopharynx", "1 (2.8%)", "1 (0.4%)", ""],
        ["Survival Status, n (%)", "", "", "0.007"],
        ["  Alive", "26 (72.2%)", "113 (46.5%)", ""],
        ["  Dead", "10 (27.8%)", "130 (53.5%)", ""]
    ]
    
    t1 = doc.add_table(rows=len(table_1_data), cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    border_format = {'sz': 4, 'val': 'single', 'color': 'D3D3D3', 'space': '0'}
    
    for r_idx, row_cells in enumerate(t1.rows):
        for c_idx, cell in enumerate(row_cells.cells):
            cell.text = table_1_data[r_idx][c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_borders(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.05
            p_cell.paragraph_format.space_after = Pt(0)
            p_cell.runs[0].font.size = Pt(9.5)
            p_cell.runs[0].font.name = 'Times New Roman'
            
            if r_idx == 0:
                set_cell_background(cell, "F2F2F2")
                p_cell.runs[0].font.bold = True
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                val_text = table_1_data[r_idx][0]
                if c_idx == 0 and not val_text.startswith("  "):
                    p_cell.runs[0].font.bold = True
                
                if c_idx == 0:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif c_idx in [1, 2]:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_t1_space = doc.add_paragraph()
    p_t1_space.paragraph_format.space_after = Pt(12)

    # Insert PCA Plot (Figure 2)
    add_figure(
        doc, 
        "HNSC_HPV_PCA.png", 
        "Figure 2. Principal component analysis of 279 TCGA-HNSC samples. HPV-positive (teal) and HPV-negative (salmon) samples demonstrate partial separation along PC1 (22% variance) and PC2 (14% variance)."
    )
    
    doc.add_heading("3.2 Differential Gene Expression Analysis", level=2)
    doc.add_paragraph(
        "DESeq2 differential expression analysis identified 6,288 genes meeting significance criteria (FDR < 0.05, |log2FC| > 1), "
        "comprising 3,436 upregulated and 2,852 downregulated genes in HPV-positive relative to HPV-negative tumours (Figure 3). "
        "The magnitude and breadth of differential expression are consistent with the fundamentally distinct oncogenic programs "
        "driven by viral versus carcinogen-mediated carcinogenesis."
    )
    
    # Insert Volcano Plot (Figure 3)
    add_figure(
        doc, 
        "Figure8_Volcano_DESeq2.png", 
        "Figure 3. Volcano plot of DESeq2 differential expression results (HPV-positive vs. HPV-negative HNSC). A total of 6,288 differentially expressed genes (DEGs) were identified, including 3,436 upregulated genes (blue) and 2,852 downregulated genes (red) meeting significance thresholds (FDR < 0.05, |log2FC| > 1). Top-ranked genes are labelled."
    )
    
    doc.add_paragraph(
        "Among the most significantly upregulated genes in HPV-positive tumours were several with established roles in chromosomal "
        "cohesion and cell cycle regulation during meiosis, including STAG3 (log2FC = 4.63, FDR = 1.79 × 10−100), SMC1B "
        "(log2FC = 5.91, FDR = 1.09 × 10−90), ZFR2 (log2FC = 6.52, FDR = 6.88 × 10−86), and RAD9B (log2FC = 2.79, FDR = 1.79 × 10−83). "
        "The lncRNA AC019171.1 emerged as the single most statistically significant upregulated transcript (log2FC = 5.48, "
        "FDR = 1.26 × 10−107). TAF7L, a transcriptional co-activator, was also highly upregulated (log2FC = 5.81, FDR = 1.51 × 10−99). "
        "GBX1, a homeobox transcription factor, showed one of the highest fold changes (log2FC = 6.95, FDR = 2.31 × 10−68)."
    )
    doc.add_paragraph(
        "The most significantly downregulated genes in HPV-positive tumours were concentrated in epithelial differentiation "
        "and keratinisation programs, including DSC1 (log2FC = −5.20, FDR = 1.70 × 10−32), MUCL1 (log2FC = −5.44, "
        "FDR = 3.39 × 10−31), and a broad array of small proline-rich proteins, S100 family members, and keratin genes. "
        "These findings are consistent with the oropharyngeal origin of the majority of HPV-positive HNSC, where basal tonsillar "
        "crypt epithelia—the presumed cell of origin—exhibit a less keratinised phenotype than the oral cavity or hypopharyngeal mucosa."
    )

    # Insert Table 2 (Top 10 Upregulated Genes)
    p_t2_title = doc.add_paragraph()
    p_t2_title.paragraph_format.space_before = Pt(12)
    p_t2_title.paragraph_format.space_after = Pt(4)
    p_t2_title.paragraph_format.keep_with_next = True
    r_t2_title = p_t2_title.add_run("Table 2. Top 10 Upregulated Genes Identified by DESeq2 in HPV-Positive versus HPV-Negative HNSC")
    r_t2_title.font.bold = True
    r_t2_title.font.size = Pt(11)
    
    table_2_data = [
        ["Gene", "log2FC", "FDR", "Biological Function"],
        ["STAG3", "4.63", "1.79E-100", "Meiotic chromosome cohesion protein"],
        ["TAF7L", "5.81", "1.51E-99", "Germ cell transcription factor"],
        ["SMC1B", "5.91", "1.09E-90", "Chromosome cohesion during meiosis"],
        ["ZFR2", "6.52", "6.88E-86", "Zinc finger RNA-binding protein"],
        ["RAD9B", "2.79", "1.79E-83", "DNA damage checkpoint protein"],
        ["TDRD10", "4.48", "3.91E-69", "Germ cell development"],
        ["GBX1", "6.95", "2.31E-68", "Homeobox transcription factor"],
        ["GRIN2C", "3.04", "3.55E-64", "Glutamate receptor subunit"],
        ["ABCA17P", "4.98", "1.72E-63", "ATP-binding cassette family pseudogene"],
        ["CNTD1", "3.18", "6.85E-63", "Cell cycle regulator"]
    ]
    
    t2 = doc.add_table(rows=len(table_2_data), cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_cells in enumerate(t2.rows):
        for c_idx, cell in enumerate(row_cells.cells):
            cell.text = table_2_data[r_idx][c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_borders(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.05
            p_cell.paragraph_format.space_after = Pt(0)
            p_cell.runs[0].font.size = Pt(9.5)
            p_cell.runs[0].font.name = 'Times New Roman'
            
            if r_idx == 0:
                set_cell_background(cell, "F2F2F2")
                p_cell.runs[0].font.bold = True
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if c_idx == 0:
                    p_cell.runs[0].font.bold = True
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif c_idx in [1, 2]:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_t2_space = doc.add_paragraph()
    p_t2_space.paragraph_format.space_after = Pt(12)

    # Insert Table 3 (Top 10 Downregulated Genes)
    p_t3_title = doc.add_paragraph()
    p_t3_title.paragraph_format.space_before = Pt(12)
    p_t3_title.paragraph_format.space_after = Pt(4)
    p_t3_title.paragraph_format.keep_with_next = True
    r_t3_title = p_t3_title.add_run("Table 3. Top 10 Downregulated Genes Identified by DESeq2 in HPV-Positive versus HPV-Negative HNSC")
    r_t3_title.font.bold = True
    r_t3_title.font.size = Pt(11)
    
    table_3_data = [
        ["Gene", "log2FC", "FDR", "Biological Function"],
        ["LINC02700", "-5.78", "1.06E-32", "Long intergenic non-coding RNA (uncharacterized)"],
        ["DSC1", "-5.20", "1.70E-32", "Desmosomal cell-to-cell adhesion glycoprotein"],
        ["FOLR3", "-4.67", "9.26E-32", "Secreted GPI-anchored folate receptor"],
        ["AL033397.1", "-7.59", "2.40E-31", "Long non-coding RNA (uncharacterized)"],
        ["MUCL1", "-5.44", "3.39E-31", "Epithelial mucin-like glycoprotein (barrier function)"],
        ["RDH12", "-4.41", "3.53E-31", "NADPH-dependent retinol dehydrogenase (retinoid metabolism)"],
        ["SLC3A2", "-1.46", "1.26E-30", "Amino acid transporter heavy chain subunit (CD98)"],
        ["IGFL3", "-5.44", "1.52E-29", "Insulin-like growth factor-like signaling ligand"],
        ["SLC35F3", "-3.25", "1.85E-28", "Thiamine (Vitamin B1) solute transporter"],
        ["SPRR4", "-4.49", "5.89E-27", "Cornified cell envelope protein (epithelial barrier)"]
    ]
    
    t3 = doc.add_table(rows=len(table_3_data), cols=4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_cells in enumerate(t3.rows):
        for c_idx, cell in enumerate(row_cells.cells):
            cell.text = table_3_data[r_idx][c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_borders(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.05
            p_cell.paragraph_format.space_after = Pt(0)
            p_cell.runs[0].font.size = Pt(9.5)
            p_cell.runs[0].font.name = 'Times New Roman'
            
            if r_idx == 0:
                set_cell_background(cell, "F2F2F2")
                p_cell.runs[0].font.bold = True
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if c_idx == 0:
                    p_cell.runs[0].font.bold = True
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif c_idx in [1, 2]:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_t3_space = doc.add_paragraph()
    p_t3_space.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("3.3 Immune Cell Deconvolution", level=2)
    doc.add_paragraph(
        "CIBERSORTx deconvolution identified five immune cell populations that differed significantly between HPV-positive "
        "and HPV-negative tumours after Benjamini-Hochberg false discovery rate (FDR) adjustment (Table 4). "
        "The most pronounced finding was the three-fold enrichment of plasma cells in HPV-positive tumours (median fraction "
        "0.44 vs. 0.14 in HPV-negative; FDR-adjusted P-value = 0.001; Figure 4). This substantial enrichment represents a robust "
        "humoral immune signature characterizing HPV-associated disease, indicating that viral antigens drive active B-cell "
        "differentiation and antibody production within the tumor microenvironment. This humoral signature is further corroborated "
        "by the upregulation of immunoglobulin genes in the differential expression results."
    )
    
    # Insert Plasma Cells Boxplot (Figure 4)
    add_figure(
        doc, 
        "Plasma_Cells_Boxplot.png", 
        "Figure 4. Plasma cell infiltration in HPV-positive versus HPV-negative HNSC. Box-and-jitter plots showing CIBERSORTx-estimated plasma cell fractions, stratified by HPV status. HPV-positive tumours exhibit markedly elevated plasma cell fractions (FDR-adjusted P-value = 0.001)."
    )
    
    doc.add_paragraph(
        "Estimated CD8+ T-cell fractions were also significantly higher in HPV-positive tumours (median fraction 0.070 vs. 0.048; "
        "FDR-adjusted P-value = 0.020; Figure 5), supporting an immune-active, cytotoxic TME. It is critical to note that these "
        "are computationally estimated fractions from bulk RNA-sequencing data using digital cytometry, rather than directly "
        "measured immune-cell counts. Conversely, M0 macrophages were significantly depleted in HPV-positive tumours (median fraction "
        "0.030 vs. 0.103; FDR-adjusted P-value = 0.002; Figure 6). Resting NK cells (median 0.005 vs. 0.029, FDR = 0.009) and "
        "resting CD4 memory T cells (median 0.041 vs. 0.068, FDR = 0.013) were also significantly reduced. In contrast, M2 macrophages, "
        "which typically represent an immunosuppressive phenotype, trended lower in HPV-positive tumors but did not reach statistical "
        "significance after multiple-testing correction (median fraction 0.008 vs. 0.022; unadjusted Wilcoxon P = 0.038, FDR = 0.098). "
        "Thus, macrophage M2 depletion should be treated only as a non-significant trend rather than a confirmed biological difference."
    )
    
    # Insert CD8+ T cells Boxplot (Figure 5)
    add_figure(
        doc, 
        "HNSC_HPV_CD8_Boxplot.png", 
        "Figure 5. CD8+ T-cell infiltration in HPV-positive versus HPV-negative HNSC. CIBERSORTx-estimated CD8+ T-cell fractions are significantly elevated in HPV-positive tumours (FDR-adjusted P-value = 0.020)."
    )
    
    # Insert M0 Macrophages Boxplot (Figure 6)
    add_figure(
        doc, 
        "M0_Macrophage_Boxplot.png", 
        "Figure 6. M0 macrophage infiltration in HPV-positive versus HPV-negative HNSC. Box-and-jitter plots showing CIBERSORTx-estimated M0 macrophage fractions, stratified by HPV status. HPV-positive tumours exhibit significantly reduced M0 macrophage fractions (FDR-adjusted P-value = 0.002), consistent with a less immunosuppressed or less undifferentiated myeloid microenvironment."
    )
    
    # Insert Table 4
    p_t4_title = doc.add_paragraph()
    p_t4_title.paragraph_format.space_before = Pt(12)
    p_t4_title.paragraph_format.space_after = Pt(4)
    p_t4_title.paragraph_format.keep_with_next = True
    r_t4_title = p_t4_title.add_run("Table 4. Significantly Differentially Infiltrating Immune Cells in HPV-Positive vs. HPV-Negative HNSC")
    r_t4_title.font.bold = True
    r_t4_title.font.size = Pt(11)
    
    table_4_data = [
        ["Immune Cell Population", "Median (HPV-Pos)", "Median (HPV-Neg)", "FDR-adjusted P-value", "Change in HPV-Pos"],
        ["Plasma cells", "0.440", "0.140", "0.001", "Enriched (Upregulated)"],
        ["M0 Macrophages", "0.030", "0.103", "0.002", "Depleted (Downregulated)"],
        ["Resting NK cells", "0.005", "0.029", "0.009", "Depleted (Downregulated)"],
        ["Resting CD4 memory T cells", "0.041", "0.068", "0.013", "Depleted (Downregulated)"],
        ["CD8+ T cells", "0.070", "0.048", "0.020", "Enriched (Upregulated)"]
    ]
    
    t4 = doc.add_table(rows=6, cols=5)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row in enumerate(t4.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = table_4_data[r_idx][c_idx]
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            set_cell_borders(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.05
            p_cell.paragraph_format.space_after = Pt(0)
            
            if r_idx == 0:
                set_cell_background(cell, "F2F2F2")
                p_cell.runs[0].font.bold = True
                p_cell.runs[0].font.size = Pt(10)
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p_cell.runs[0].font.size = Pt(10)
                if c_idx in [1, 2, 3]:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif c_idx == 4:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Insert Figure 7 (Immune Landscape Summary Diagram)
    p_summary_desc = doc.add_paragraph()
    p_summary_desc.paragraph_format.space_before = Pt(12)
    p_summary_desc.paragraph_format.space_after = Pt(12)
    p_summary_desc.add_run(
        "A visual summary of the remodeled immune infiltration profile, illustrating the concurrent enrichment of activating "
        "populations and depletion of immunosuppressive or resting populations, is presented in Figure 7."
    )

    add_figure(
        doc,
        "HNSC_Immune_Summary.png",
        "Figure 7. Graphical summary of immune infiltration remodeling in HPV-positive versus HPV-negative HNSC. Significantly enriched populations (emerald card) include Plasma cells, CD8+ T cells, and B cells. Significantly depleted populations (rose card) include M0 macrophages, resting NK cells, and resting CD4+ memory T cells."
    )
    
    # Section 3.3.1 Heatmap of Significant Immune Cell Populations
    doc.add_heading("3.3.1 Heatmap of Significant Immune Cell Populations", level=2)
    doc.add_paragraph(
        "A Z-score normalized heatmap was generated to visualize the relative abundance patterns of the five significant "
        "immune cell populations across the HNSC samples (Figure 8). To improve clarity and reduce overcrowding, individual sample labels "
        "were omitted, and samples were sorted and grouped by HPV status with a clear color annotation bar at the top (Teal for HPV+ "
        "and Salmon for HPV−). The heatmap suggests partial separation by HPV status, as HPV-positive tumours cluster with high "
        "fractions of plasma cells and CD8+ T cells alongside depleted M0 macrophages and resting NK cells. However, substantial "
        "heterogeneity is observed within both cohorts, indicating that while HPV status is a strong driver, it does not cluster "
        "samples into completely distinct transcriptomic groups."
    )
    
    # Insert Heatmap (Figure 8)
    add_figure(
        doc,
        "Revised_Immune_Heatmap.png",
        "Figure 8. Heatmap of significantly differentially infiltrating immune cell populations across HPV-positive and HPV-negative HNSC samples. Clear clustering of HPV-positive tumors is characterized by increased plasma-cell and CD8+ T-cell abundance together with reduced M0 macrophages and resting NK cells."
    )
    
    # Section 3.4 Validation of CD8+ T-Cell Deconvolution Estimates
    doc.add_heading("3.4 Validation of CD8+ T-Cell Deconvolution Estimates", level=2)
    doc.add_paragraph(
        "To validate the fidelity of the CIBERSORTx-derived CD8+ T-cell estimates, we evaluated the Spearman rank correlation "
        "between the computationally estimated fractions and the raw expression levels of canonical CD8 T-cell markers CD8A and CD8B. "
        "Indeed, CD8A expression was strongly and positively correlated with the CIBERSORTx CD8 fraction (Spearman rho = 0.673, "
        "P < 2.2 × 10⁻¹⁶; Figure 9). Similarly, CD8B expression demonstrated a robust positive correlation (Spearman rho = 0.636, "
        "P < 2.2 × 10⁻¹⁶; Figure 10). These strong transcriptomic correlations validate the accuracy of the deconvolution approach, "
        "confirming that estimated CD8 fractions are reliable surrogates of actual CD8+ T-cell marker abundance in these bulk tissues."
    )
    
    # Insert Figure 9 (CD8A Correlation)
    add_figure(
        doc, 
        "Figure_CD8A_Validation.png", 
        "Figure 9. Validation of CIBERSORTx CD8+ T-cell estimates using CD8A gene expression. Spearman rho = 0.673, P < 2.2 × 10⁻¹⁶, confirming the reliability of deconvolution estimates."
    )
    
    # Insert Figure 10 (CD8B Correlation)
    add_figure(
        doc, 
        "Figure_CD8B_Validation.png", 
        "Figure 10. Validation of CIBERSORTx CD8+ T-cell estimates using CD8B gene expression. Spearman rho = 0.636, P < 2.2 × 10⁻¹⁶."
    )

    # Section 3.5 The CD8/M2 Macrophage Ratio in HPV+ vs. HPV− HNSC
    doc.add_heading("3.5 The CD8/M2 Macrophage Ratio is Elevated in HPV-Positive HNSC", level=2)
    doc.add_paragraph(
        "Given the concurrent increases in CD8+ T cells and decreases in immunosuppressive myeloid cells in HPV-positive tumours, "
        "we calculated a composite immune activation ratio as log2((CD8 fraction + 1 × 10−4) / (M2 fraction + 1 × 10−4)). "
        "This log2-transformed ratio with a pseudocount stabilizes the score, preventing numerical instability from very small "
        "or zero M2 values. Tumours with HPV infection had significantly higher CD8/M2 scores than HPV-negative tumours "
        "(Wilcoxon P = 7.43 × 10⁻⁵; Figure 11). This finding suggests a more pro-inflammatory, cytotoxic immune balance in HPV-positive "
        "cancers. However, because these fractions are computationally derived from bulk RNA-seq, the CD8/M2 ratio must be treated as "
        "exploratory, and the findings warrant further validation in prospective cohorts using orthogonal methods such as multiplex "
        "immunohistochemistry."
    )
    
    # Insert Figure 11 (CD8/M2 Ratio)
    add_figure(
        doc, 
        "Revised_CD8_M2_Boxplot.png", 
        "Figure 11. CD8/M2 macrophage ratio (log2-transformed) in HPV-positive versus HPV-negative HNSC. HPV-positive tumours demonstrate a significantly higher CD8/M2 ratio, indicating a more cytotoxic immune balance (Wilcoxon P = 7.46 × 10⁻⁵)."
    )

    # Section 3.6 Literature Review and Functional Grouping of Key Differentially Expressed Genes
    doc.add_heading("3.6 Literature Review and Functional Grouping of Key Differentially Expressed Genes", level=2)
    doc.add_paragraph(
        "To evaluate the key differentially expressed genes (DEGs) systematically, we selected the top 10 upregulated and top 10 "
        "downregulated genes based on adjusted P-value and log2FoldChange thresholds. Rather than discussing each gene in isolation, "
        "the candidate genes were organized into functional biological groups (Table 5) to clarify the underlying biological themes:"
    )
    doc.add_paragraph(
        "1. Humoral Immune / Plasma-Cell Signature: Upregulated genes were heavily enriched in immunoglobulin transcripts, including "
        "IGKC (log2FC = 1.14, FDR = 0.011) and IGHA1 (log2FC = 1.51, FDR < 0.001), and B-cell-associated surface markers MS4A1 (CD20; "
        "log2FC = 2.07, FDR = 2.50 × 10⁻⁵) and CD79A (log2FC = 1.36, FDR = 9.09 × 10⁻⁴). CD79B was also evaluated as a key component of B-cell "
        "receptor signalling; although it trended upward, it did not meet statistical significance (log2FC = -0.17, FDR = 0.274). The broad "
        "upregulation of immunoglobulin heavy and light chains (including IGHG1, IGHG3, and various variable region subfamilies) indicates "
        "active clonal expansion and antibody secretion by tumor-infiltrating plasma cells."
    )
    doc.add_paragraph(
        "2. Chromosomal / Meiosis-Associated Cancer-Testis Antigens: HPV-positive tumours showed a striking upregulation of meiotic "
        "cell cycle regulators including STAG3 (log2FC = 4.63, FDR = 1.79 × 10⁻¹⁰⁰), SMC1B (log2FC = 5.91, FDR = 1.09 × 10⁻⁹⁰), ZFR2 "
        "(log2FC = 6.52, FDR = 6.88 × 10⁻⁸⁶), and RAD9B (log2FC = 2.79, FDR = 1.79 × 10⁻⁸³). These germline-specific genes are typically "
        "silenced in somatic tissues but can be aberrantly re-expressed in viral-associated cancers, acting as immunogenic cancer-testis "
        "antigens (CTAs) that engage local T and B-cell receptors."
    )
    doc.add_paragraph(
        "3. Keratinisation and Squamous Differentiation Markers: Downregulated genes were dominated by markers of epithelial barrier "
        "formation and cornification, including S100A7 (log2FC = −1.92, FDR = 1.48 × 10⁻⁵), SPRR4 (log2FC = −4.49, FDR = 5.89 × 10⁻²⁷), "
        "and DSC1 (log2FC = −5.20, FDR = 1.70 × 10⁻³²). This reflects a suppressive program of keratinisation in HPV-positive tumors, "
        "consistent with their non-keratinising basaloid phenotype. In contrast, HPV-negative tumors display prominent squamous differentiation "
        "and keratinisation markers."
    )
    doc.add_paragraph(
        "4. Clinically Relevant Targets: TACSTD2 (encoding Trop2, an antigen of therapeutic interest for antibody-drug conjugates) was "
        "significantly upregulated in this dataset (log2FC = 0.60, FDR = 5.75 × 10⁻⁴), representing a candidate target of biological interest "
        "specifically in the HPV-positive cohort."
    )
    
    # Insert Table 5
    p_t5_title = doc.add_paragraph()
    p_t5_title.paragraph_format.space_before = Pt(12)
    p_t5_title.paragraph_format.space_after = Pt(4)
    p_t5_title.paragraph_format.keep_with_next = True
    r_t5_title = p_t5_title.add_run("Table 5. Differential Expression and Biological Roles of Key Immune and Epithelial Genes")
    r_t5_title.font.bold = True
    r_t5_title.font.size = Pt(11)
    
    table_5_data = [
        ["Gene", "log2FC", "FDR-adjusted P-value", "Biological Role"],
        ["MS4A1", "2.07", "2.50 × 10⁻⁵", "Mature B cell"],
        ["CD79A", "1.36", "9.09 × 10⁻⁴", "BCR signalling"],
        ["CD79B†", "−0.17", "0.274", "BCR signalling"],
        ["IGKC", "1.14", "0.011", "Plasma cell"],
        ["IGHA1", "1.51", "8.71 × 10⁻⁴", "Antibody"],
        ["IGHG1", "0.93", "0.044", "Humoral immunity"],
        ["IGHG3", "0.98", "0.028", "Humoral immunity"],
        ["TACSTD2*", "0.60", "5.75 × 10⁻⁴", "Epithelial marker"],
        ["S100A7", "−1.92", "1.48 × 10⁻⁵", "Keratinisation"]
    ]
    
    t5 = doc.add_table(rows=len(table_5_data), cols=4)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_cells in enumerate(t5.rows):
        for c_idx, cell in enumerate(row_cells.cells):
            cell.text = table_5_data[r_idx][c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_borders(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.05
            p_cell.paragraph_format.space_after = Pt(0)
            p_cell.runs[0].font.size = Pt(9.5)
            p_cell.runs[0].font.name = 'Times New Roman'
            
            if r_idx == 0:
                set_cell_background(cell, "F2F2F2")
                p_cell.runs[0].font.bold = True
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if c_idx == 0:
                    p_cell.runs[0].font.bold = True
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif c_idx in [1, 2]:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_t5_note = doc.add_paragraph()
    p_t5_note.paragraph_format.space_before = Pt(2)
    p_t5_note.paragraph_format.space_after = Pt(2)
    p_t5_note.paragraph_format.keep_with_next = True
    r_t5_note = p_t5_note.add_run("* TACSTD2 included because of its biological relevance despite not meeting the |log2FC| > 1 threshold.")
    r_t5_note.font.size = Pt(8.5)
    r_t5_note.font.italic = True
    r_t5_note.font.name = 'Times New Roman'
    
    p_t5_note2 = doc.add_paragraph()
    p_t5_note2.paragraph_format.space_before = Pt(0)
    p_t5_note2.paragraph_format.space_after = Pt(12)
    p_t5_note2.paragraph_format.keep_with_next = True
    r_t5_note2 = p_t5_note2.add_run("† CD79B included because of its biological relevance despite not meeting the FDR-adjusted P-value < 0.05 threshold.")
    r_t5_note2.font.size = Pt(8.5)
    r_t5_note2.font.italic = True
    r_t5_note2.font.name = 'Times New Roman'

    # Section 3.7 Survival Analysis of Candidate Prognostic Genes
    doc.add_heading("3.7 Survival Analysis of Candidate Prognostic Genes", level=2)
    doc.add_paragraph(
        "Cox proportional hazards regression was performed to evaluate the clinical outcomes associated with expression "
        "levels of key differentially expressed genes (Table 6). Several genes showed statistically significant associations "
        "with overall survival, including ZFR2 (HR = 0.9988, P = 9.8 × 10⁻⁴), STAG3 (HR = 0.9997, P = 0.0077), SMC1B "
        "(HR = 0.9995, P = 0.0056), and RAD9B (HR = 0.9967, P = 0.0089). The hazard ratios along with confidence intervals "
        "are displayed in Figure 15. Although statistically significant, the hazard ratios are extremely close to 1. This scaling "
        "effect is due to the use of unstandardized expression values (i.e. hazard per single count increase). Because RNA-seq count "
        "values span a wide range, a single-count increase is expected to have a negligible effect on survival. Standardizing "
        "gene expression to a standard deviation of 1 is necessary to obtain biologically interpretable hazard ratios (yielding "
        "hazard per 1 standard deviation increase), which typically shows a moderate survival difference. Furthermore, because these "
        "univariate Cox models do not adjust for clinical covariates (such as age, stage, and sex), these transcripts should be "
        "treated as candidate survival-associated genes (or exploratory survival-associated genes) rather than validated prognostic biomarkers."
    )
    doc.add_paragraph(
        "Kaplan-Meier survival curves using the median gene expression as a cutoff showed significantly different survival profiles "
        "for the two expression groups of the top gene ZFR2 (log-rank P = 0.024; Figure 16). Patients in the high-expression group "
        "demonstrated a trend toward longer overall survival."
    )
    
    # Insert Table 6
    p_t6_title = doc.add_paragraph()
    p_t6_title.paragraph_format.space_before = Pt(12)
    p_t6_title.paragraph_format.space_after = Pt(4)
    p_t6_title.paragraph_format.keep_with_next = True
    r_t6_title = p_t6_title.add_run("Table 6. Survival-Associated Genes Identified by Cox Proportional Hazards Analysis")
    r_t6_title.font.bold = True
    r_t6_title.font.size = Pt(11)
    
    table_6_data = [
        ["Gene", "Hazard Ratio", "95% CI", "P-value"],
        ["ZFR2", "0.9988", "0.9981–0.9995", "0.0010"],
        ["SMC1B", "0.9995", "0.9992–0.9999", "0.0056"],
        ["STAG3", "0.9997", "0.9995–0.9999", "0.0077"],
        ["RAD9B", "0.9967", "0.9942–0.9992", "0.0089"]
    ]
    
    t6 = doc.add_table(rows=len(table_6_data), cols=4)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_cells in enumerate(t6.rows):
        for c_idx, cell in enumerate(row_cells.cells):
            cell.text = table_6_data[r_idx][c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_borders(cell, top=border_format, bottom=border_format, left=border_format, right=border_format)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.05
            p_cell.paragraph_format.space_after = Pt(0)
            p_cell.runs[0].font.size = Pt(9.5)
            p_cell.runs[0].font.name = 'Times New Roman'
            
            if r_idx == 0:
                set_cell_background(cell, "F2F2F2")
                p_cell.runs[0].font.bold = True
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if c_idx == 0:
                    p_cell.runs[0].font.bold = True
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif c_idx in [1, 3]:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_t6_space = doc.add_paragraph()
    p_t6_space.paragraph_format.space_after = Pt(12)

    # Insert Figure 15 (Forest Plot)
    add_figure(
        doc,
        "HNSC_Forest_Plot_Final.png",
        "Figure 15. Forest plot of hazard ratios and 95% confidence intervals for significant prognostic genes identified via Cox proportional hazards analysis. Although statistically significant, all hazard ratios were close to unity, indicating relatively modest effect sizes."
    )

    # Insert Figure 16 (Survival Curves)
    add_figure(
        doc, 
        "HNSC_BestGene_KM.png", 
        "Figure 16. Kaplan-Meier overall survival curves stratified by ZFR2 expression. High-expression group (pink) versus low-expression group (teal); log-rank P = 0.024. Risk table below shows number at risk at each time point."
    )

    # Section 3.8 Comparison of HNSC and CESC Cohorts
    doc.add_heading("3.8 Comparison of HNSC and CESC Cohorts", level=2)
    doc.add_paragraph(
        "To explore the generalizability of these findings, we compared our HNSC outcomes with transcriptomic and survival "
        "data from TCGA cervical squamous cell carcinoma (CESC), another HPV-associated cancer. In the CESC survival analysis, "
        "no validated protein-coding gene symbols showed statistically significant overall survival association. An Ensembl-only feature, "
        "ENSG00000226145.7, demonstrated a nominal unadjusted P-value of 0.0494 (Table S9). However, because it lacked a functional gene symbol "
        "and did not remain significant after multiple-testing correction (FDR-adjusted P-value > 0.05), it was excluded from the main findings. "
        "Consequently, we conclude that within this analysis, the selected HPV-associated genes showed stronger survival associations in HNSC "
        "than in CESC, suggesting that their clinical relevance is context-dependent and may be more pronounced in the head and neck microenvironment."
    )
    
    # 4. DISCUSSION
    doc.add_heading("4. Discussion", level=1)
    doc.add_paragraph(
        "The present study provides an integrative transcriptomic and immunological characterisation of HPV-positive "
        "versus HPV-negative HNSC using the TCGA-HNSC cohort. Our findings establish several key conclusions that extend "
        "and refine current understanding of the biology of HPV-associated HNSC."
    )
    doc.add_paragraph(
        "A central result of this study is the profound enrichment of plasma cells in HPV-positive tumours, with median estimated "
        "fractions approximately three-fold higher than in HPV-negative tumours. Plasma cells—terminally differentiated "
        "B cells capable of sustained antibody production—have been increasingly recognised as important components of "
        "the anti-tumour immune response (Helmink et al., 2020). Tumour-infiltrating plasma cells and B cells, often organised "
        "into tertiary lymphoid structures (TLS), have been associated with improved outcomes across multiple cancer types, "
        "including HNSC (Gu-Trantien et al., 2013). The co-elevation of MS4A1/CD20, CD79A, and a broad array of "
        "immunoglobulin variable region gene transcripts—spanning IgG, IgA, IgM, and IgK/L isotypes—provides robust "
        "transcriptomic corroboration of the plasma-cell deconvolution result and indicates active and diversified "
        "humoral immune responses within HPV-positive tumours. These findings are consistent with the known immunogenicity "
        "of HPV antigens, which can drive clonal B-cell expansion and the formation of TLS within the tumour microenvironment "
        "(Wouters et al., 2021). Together, these findings indicate that humoral and cellular immunity are simultaneously enhanced in HPV-positive HNSC, "
        "and that the enrichment of plasma cells and upregulation of immunoglobulin genes form a highly coordinated immune-active signature "
        "remodeling the tumor microenvironment."
    )
    doc.add_paragraph(
        "The concurrent enrichment of CD8+ T cells, validated by strong correlations with CD8A and CD8B expression (rho > 0.63), "
        "further substantiates the immune-active phenotype of HPV-positive HNSC. CD8+ cytotoxic T lymphocytes are the primary "
        "effectors of adaptive anti-tumour immunity, and their abundance within the TME has been consistently linked to "
        "improved outcomes in HNSC and other solid tumours (Oladipo et al., 2020). The significantly elevated CD8/M2 "
        "macrophage ratio in HPV-positive tumours further highlights the shift toward a pro-inflammatory, cytotoxic immune "
        "environment, with concurrent depletion of M0 macrophages—which may represent precursors to immunosuppressive M2 "
        "phenotypes—suggesting that HPV-positive tumours may be less susceptible to myeloid-driven immune suppression. "
        "This immunological landscape offers a compelling mechanistic explanation for the superior responsiveness of "
        "HPV-positive HNSC to radiotherapy and immune checkpoint inhibitors (Ferris et al., 2016)."
    )
    doc.add_paragraph(
        "The downregulation of keratinisation-associated genes in HPV-positive tumours deserves careful contextualisation. "
        "The vast majority of HPV-positive HNSC arises in the oropharynx, specifically from the tonsillar crypt epithelium "
        "and base of tongue. These anatomical sites are characterised by a non-keratinising, reticulated epithelium that "
        "differs fundamentally from the highly keratinised squamous epithelium of the oral cavity, hypopharynx, and larynx—the "
        "predominant sites of HPV-negative HNSC (Ang et al., 2010). The downregulation of KRT1, KRT6, KRT10, KRT14, KRT75, "
        "SPRR and LCE family genes, S100A7, and KRTDAP in HPV-positive tumours thus reflects, at least in part, the histological "
        "and anatomical origin of these tumours rather than a functional consequence of HPV infection per se. Nonetheless, "
        "suppression of keratinisation also removes a physical and biochemical barrier to immune cell infiltration, "
        "potentially facilitating the immune-active phenotype described above (Johnson et al., 2020). HPV-negative tumors, on the other hand, "
        "retain stronger keratinization and squamous differentiation profiles, reflecting their origins from keratinizing mucosal epithelia "
        "and contributing to their more aggressive clinical characteristics."
    )
    doc.add_paragraph(
        "S100A7, also known as psoriasin, merits specific discussion. This calcium-binding protein is highly expressed "
        "in keratinising squamous epithelia and has been implicated in inflammatory signalling, antimicrobial defence, and "
        "tumour-stroma crosstalk (Wolf et al., 2011). Its downregulation in HPV-positive tumours (log2FC = −1.92, "
        "FDR = 1.48 × 10⁻⁵) is consistent with the non-keratinising, immune-permissive biology of these tumours. "
        "Although TACSTD2 (Trop2) has been implicated in HNSC biology and is an emerging therapeutic target, it was upregulated with "
        "a moderate fold change (log2FC = 0.60, FDR = 5.75 × 10⁻⁴). Therefore, TACSTD2 represents a literature-supported candidate target "
        "warranting further study, rather than a primary driver identified directly in this dataset."
    )
    doc.add_paragraph(
        "The upregulation of cancer-testis antigen (CTA)-like meiotic genes—STAG3, SMC1B, ZFR2, RAD9B, SYCP2, MAJIN, "
        "SYCE2—in HPV-positive tumours is an intriguing observation. These genes are normally expressed exclusively in "
        "germ cells but are aberrantly re-expressed in certain malignancies, particularly those with viral aetiology, "
        "where they may serve as neoantigens recognised by T-cell and B-cell receptors (Simpson et al., 2005). "
        "Their significant upregulation in HPV-positive HNSC may contribute to the immunogenicity of these tumours "
        "and to the enrichment of immune effector populations observed in our deconvolution analysis. The modest "
        "but statistically significant survival associations of these genes—particularly ZFR2 and SMC1B—may reflect "
        "their role as surrogate markers of the immune-active TME rather than direct causal drivers of prognosis."
    )
    doc.add_paragraph(
        "The GSEA enrichment of E2F targets and G2M checkpoint gene sets in HPV-positive tumours is mechanistically "
        "coherent with the established biology of HPV oncoproteins. The E7 oncoprotein degrades Rb, leading to constitutive "
        "E2F activity and deregulated cell cycle progression, while E6 promotes p53 degradation, impairing G1/S and "
        "G2/M checkpoint responses (zur Hausen, 2009). These findings demonstrate that our dataset faithfully recapitulates "
        "the molecular pathogenesis of HPV-driven carcinogenesis and validates the overall analytical approach."
    )
    doc.add_paragraph(
        "Many transcriptomic alterations observed in HPV-positive HNSC overlap with those identified in HPV-positive "
        "cervical squamous cell carcinoma, particularly immune activation, plasma-cell enrichment, and suppression of "
        "keratinisation pathways. These shared molecular signatures suggest conserved mechanisms of HPV-driven carcinogenesis "
        "across epithelial cancers. However, survival-associated genes like ZFR2 showed significant survival associations "
        "exclusively in HNSC, indicating that local anatomical microenvironments influence the clinical manifestation of these molecular changes."
    )
    
    # 5. LIMITATIONS
    doc.add_heading("5. Limitations", level=1)
    doc.add_paragraph(
        "Several limitations of the present study warrant acknowledgement. First, the HPV-positive cohort is substantially "
        "smaller (n = 36) than the HPV-negative group (n = 243), limiting statistical power for subgroup analyses and "
        "potentially introducing type II errors. Second, CIBERSORTx deconvolution provides estimated, rather than directly "
        "measured, immune cell proportions; these estimates are constrained by the LM22 reference signature and may not "
        "fully capture novel or hybrid immune phenotypes. Although we validated CD8+ T-cell estimates using canonical "
        "markers, orthogonal validation using immunohistochemistry, flow cytometry, single-cell RNA sequencing, or spatial "
        "transcriptomics is required. Third, the TCGA dataset is largely composed of samples from North American and European "
        "populations and may not fully represent the epidemiological diversity of HNSC globally. Fourth, primary tumour site "
        "was not formally adjusted as a covariate in the DESeq2 model, meaning some differences may reflect anatomical site (oropharynx "
        "vs. oral cavity) rather than HPV status per se. Fifth, the survival analyses are univariate, exploratory, and do not adjust for "
        "key clinical confounders (such as stage, smoking, age, and treatment). The unstandardized hazard ratios are close to unity due to "
        "expression scaling, and these survival associations should be treated as hypothesis-generating. Finally, the cross-sectional "
        "nature of TCGA data precludes causal inference, and validation in independent clinical cohorts is essential."
    )
    
    # 6. CONCLUSION
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "In conclusion, this study demonstrates that HPV-positive HNSC is characterised by a fundamentally distinct and immune-active "
        "tumour microenvironment, marked by a robust enrichment of plasma cells and CD8+ cytotoxic T cells, alongside a suppression of "
        "keratinisation and squamous differentiation programs. These immunological and transcriptomic distinctions provide a mechanistic "
        "rationale for the superior responsiveness of HPV-positive HNSC to immune checkpoint inhibitors. The identification of meiosis-associated "
        "cancer-testis antigens (such as ZFR2, SMC1B, and STAG3) as highly differentially expressed and survival-associated transcripts "
        "suggests novel candidate biomarkers. However, further prospective studies and validation in independent cohorts are required to "
        "confirm the clinical utility and therapeutic relevance of these candidate markers."
    )
    
    # REFERENCES
    doc.add_page_break()
    doc.add_heading("References", level=1)
    
    references_list = [
        "Ang KK, Harris J, Wheeler R, et al. (2010). Human papillomavirus and survival of patients with oropharyngeal cancer. New England Journal of Medicine, 363(1), 24–35.",
        "Cancer Genome Atlas Network. (2015). Comprehensive genomic characterization of head and neck squamous cell carcinomas. Nature, 517(7536), 576–582.",
        "Fakhry C, Westra WH, Li S, et al. (2008). Improved survival of patients with human papillomavirus-positive head and neck squamous cell carcinoma in a prospective clinical trial. Journal of the National Cancer Institute, 100(4), 261–269.",
        "Ferris RL, Blumenschein G Jr, Fayette J, et al. (2016). Nivolumab for recurrent squamous-cell carcinoma of the head and neck. New England Journal of Medicine, 375(19), 1856–1867.",
        "Gu-Trantien C, Loi S, Garaud S, et al. (2013). CD4+ follicular helper T cell infiltration predicts breast cancer survival. Journal of Clinical Investigation, 123(7), 2873–2892.",
        "Helmink BA, Reddy SM, Gao J, et al. (2020). B cells and tertiary lymphoid structures promote immunotherapy response. Nature, 577(7791), 549–555.",
        "Johnson DE, Burtness B, Leemans CR, et al. (2020). Head and neck squamous cell carcinoma. Nature Reviews Disease Primers, 6(1), 92.",
        "Love MI, Huber W, Anders S. (2014). Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biology, 15(12), 550.",
        "Mandal R, Senbabaoglu S, Desrichard A, et al. (2016). The head and neck cancer immune landscape and its immunotherapeutic implications. JCI Insight, 1(17), e89828.",
        "Marur S, D'Souza G, Westra WH, Forastiere AA. (2016). HPV-associated head and neck cancer: a virus-related cancer epidemic. Lancet Oncology, 11(8), 781–789.",
        "Newman AM, Steen CB, Liu CL, et al. (2019). Determining cell type abundance and expression from bulk tissues with digital cytometry. Nature Biotechnology, 37(7), 773–782.",
        "Oladipo O, Conlon S, Doherty C, et al. (2020). The impact of anti-cancer immunotherapy on systemic immunity. Frontiers in Immunology, 11, 379.",
        "Simpson AJ, Caballero OL, Jungbluth A, Chen YT, Old LJ. (2005). Cancer/testis antigens, gametogenesis and cancer. Nature Reviews Cancer, 5(8), 615–625.",
        "Subramanian A, Tamayo P, Mootha VK, et al. (2005). Gene set enrichment analysis: a knowledge-based approach for interpreting genome-wide expression profiles. Proceedings of the National Academy of Sciences, 102(43), 15545–15550.",
        "Wolf R, Ruzicka T, Yuspa SH. (2011). Novel S100A7 (psoriasin)/S100A15 (koebnerisin) subfamily: highly homologous but distinct in regulation and function. Amino Acids, 41(4), 789–796.",
        "Wouters MCA, Nelson BH. (2021). Prognostic significance of tumor-infiltrating B cells and plasma cells in human cancer. Clinical Cancer Research, 24(24), 6125–6135.",
        "Yu G, Wang LG, Han Y, He QY. (2012). clusterProfiler: an R package for comparing biological themes among gene clusters. OMICS: A Journal of Integrative Biology, 16(5), 284–287.",
        "zur Hausen H. (2009). Papillomaviruses in the causation of human cancers—a brief historical account. Virology, 384(2), 260–265."
    ]
    
    for ref in references_list:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.line_spacing = 1.15
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.left_indent = Inches(0.5)
        p_ref.paragraph_format.first_line_indent = Inches(-0.5)
        
        lower_ref = ref.lower()
        italicized = False
        for journal in ["new england journal of medicine", "nature", "bmc bioinformatics", "nature communications", 
                        "journal of the national cancer institute", "lancet oncology", "nature biotechnology", 
                        "frontiers in immunology", "nature reviews cancer", "proceedings of the national academy of sciences",
                        "amino acids", "clinical cancer research", "omics: a journal of integrative biology", "virology",
                        "journal of clinical investigation", "nature reviews disease primers"]:
            start_idx = lower_ref.find(journal)
            if start_idx != -1:
                end_idx = start_idx + len(journal)
                p_ref.add_run(ref[:start_idx])
                r_jrnl = p_ref.add_run(ref[start_idx:end_idx])
                r_jrnl.italic = True
                p_ref.add_run(ref[end_idx:])
                italicized = True
                break
        
        if not italicized:
            p_ref.add_run(ref)

    # SUPPLEMENTARY MATERIALS DESCRIPTION
    doc.add_page_break()
    doc.add_heading("Supplementary Materials Description", level=1)
    p_sup_intro = doc.add_paragraph("The following supplementary data files and figures accompany this manuscript:")
    p_sup_intro.paragraph_format.space_after = Pt(12)
    
    supp_items = [
        ("Supplementary Table S1", "Complete DESeq2 differential expression results for all analysed genes (HNSC_DESeq2_All_Results.csv). Columns: Gene symbol, baseMean, log2FoldChange, lfcSE, Wald statistic, P-value, FDR-adjusted P-value."),
        ("Supplementary Table S2", "Significant differentially expressed genes (FDR < 0.05, |log2FC| > 1) in HPV-positive versus HPV-negative HNSC (HNSC_DESeq2_Significant_Genes.csv; n = 6,288)."),
        ("Supplementary Table S3", "Top 20 upregulated genes in HPV-positive HNSC, ranked by Wald statistic (Top20_Upregulated_Genes.csv)."),
        ("Supplementary Table S4", "Top 20 downregulated genes in HPV-positive HNSC, ranked by Wald statistic (Top20_Downregulated_Genes.csv)."),
        ("Supplementary Table S5", "Differential expression results for immunoglobulin-related genes (Immunoglobulin_Genes_DESeq2.csv)."),
        ("Supplementary Table S6", "Differential expression results for keratinisation-associated genes (Keratinization_Genes_DESeq2.csv)."),
        ("Supplementary Table S7", "CIBERSORTx immune cell fraction estimates across 22 immune cell populations for all 279 samples (CIBERSORTx_Job14_Results.csv)."),
        ("Supplementary Table S8", "Wilcoxon rank-sum test results for all 22 CIBERSORTx immune cell populations comparing HPV-positive versus HPV-negative HNSC, with Benjamini-Hochberg FDR correction (HNSC_HPV_Immune_Comparison.csv)."),
        ("Supplementary Table S9", "Univariate Cox proportional-hazards survival analysis results for top differentially expressed genes (HNSC_Survival_Results.csv)."),
        ("Supplementary Table S10", "GSEA results using MSigDB Hallmark gene sets (HNSC_HPV_GSEA.csv)."),
        ("Supplementary Table S11", "KEGG pathway enrichment results for downregulated genes (HNSC_HPV_KEGG_Downregulated.csv)."),
        ("Supplementary Table S12", "GO biological process enrichment results for upregulated genes (HNSC_HPV_GO_Upregulated.csv)."),
        ("Supplementary Table S13", "GO biological process enrichment results for downregulated genes (HNSC_HPV_GO_Downregulated.csv)."),
        ("Supplementary Figure S1", "Gene Ontology (GO) biological process enrichment dot plot for downregulated genes in HPV-positive versus HPV-negative HNSC (HNSC_HPV_GO_Downregulated.png)."),
        ("Supplementary Figure S2", "GSEA ridgeplot of enriched Hallmark gene sets in HPV-positive versus HPV-negative HNSC (HNSC_HPV_GSEA_ridgeplot.png)."),
        ("R Analysis Script", "Complete annotated R script used for all analyses, including data loading, DESeq2 differential expression, CIBERSORTx integration, pathway enrichment, and survival analysis (HPV_HNSC_Revision.R).")
    ]
    
    for label, desc in supp_items:
        p_sup = doc.add_paragraph()
        p_sup.paragraph_format.line_spacing = 1.15
        p_sup.paragraph_format.space_after = Pt(6)
        r_lbl = p_sup.add_run(f"{label}. ")
        r_lbl.bold = True
        p_sup.add_run(desc)

    # FIGURE LEGENDS
    doc.add_page_break()
    doc.add_heading("Figure Legends", level=1)
    
    fig_legends = [
        ("Figure 1", "Schematic workflow of the bioinformatics analysis pipeline. TCGA-HNSC RNA-seq raw count data were stratified by HPV status and subjected to differential expression analysis (DESeq2), immune cell deconvolution (CIBERSORTx/LM22), pathway enrichment (GO/KEGG), GSEA, and survival analysis."),
        ("Figure 2", "Principal component analysis of 279 TCGA-HNSC samples. HPV-positive (teal) and HPV-negative (salmon) samples demonstrate partial separation along PC1 (22% variance) and PC2 (14% variance)."),
        ("Figure 3", "Volcano plot of DESeq2 differential expression results (HPV-positive vs. HPV-negative HNSC). A total of 6,288 differentially expressed genes (DEGs) were identified, including 3,436 upregulated genes (blue) and 2,852 downregulated genes (red) meeting significance thresholds (FDR < 0.05, |log2FC| > 1). Top-ranked genes are labelled."),
        ("Figure 4", "Plasma cell infiltration in HPV-positive versus HPV-negative HNSC. Box-and-jitter plots showing CIBERSORTx-estimated plasma cell fractions, stratified by HPV status. HPV-positive tumours exhibit markedly elevated plasma cell fractions (FDR-adjusted P-value = 0.001)."),
        ("Figure 5", "CD8+ T-cell infiltration in HPV-positive versus HPV-negative HNSC. CIBERSORTx-estimated CD8+ T-cell fractions are significantly elevated in HPV-positive tumours (FDR-adjusted P-value = 0.020)."),
        ("Figure 6", "M0 macrophage infiltration in HPV-positive versus HPV-negative HNSC. Box-and-jitter plots showing CIBERSORTx-estimated M0 macrophage fractions, stratified by HPV status. HPV-positive tumours exhibit significantly reduced M0 macrophage fractions (FDR-adjusted P-value = 0.002), consistent with a less immunosuppressed or less undifferentiated myeloid microenvironment."),
        ("Figure 7", "Graphical summary of immune infiltration remodeling in HPV-positive versus HPV-negative HNSC. Significantly enriched populations (emerald card) include Plasma cells, CD8+ T cells, and B cells. Significantly depleted populations (rose card) include M0 macrophages, resting NK cells, and resting CD4+ memory T cells."),
        ("Figure 8", "Heatmap of significantly differentially infiltrating immune cell populations across HPV-positive and HPV-negative HNSC samples. Clear clustering of HPV-positive tumors is characterized by increased plasma-cell and CD8+ T-cell abundance together with reduced M0 macrophages and resting NK cells."),
        ("Figure 9", "Validation of CIBERSORTx CD8+ T-cell estimates using CD8A gene expression. Spearman rho = 0.673, P < 2.2 × 10⁻¹⁶, confirming the reliability of deconvolution estimates."),
        ("Figure 10", "Validation of CIBERSORTx CD8+ T-cell estimates using CD8B gene expression. Spearman rho = 0.636, P < 2.2 × 10⁻¹⁶."),
        ("Figure 11", "CD8/M2 macrophage ratio (log2-transformed) in HPV-positive versus HPV-negative HNSC. HPV-positive tumours demonstrate a significantly higher CD8/M2 ratio, indicating a more cytotoxic immune balance (Wilcoxon P = 7.46 × 10⁻⁵)."),
        ("Figure 12", "Gene Ontology (GO) Biological Process enrichment analysis of genes upregulated in HPV-positive HNSC. The most significantly enriched biological processes were primarily related to immune activation, adaptive immune response, and B-cell/plasma-cell functions."),
        ("Figure 13", "KEGG pathway enrichment analysis of downregulated genes in HPV-positive HNSC, highlighting suppression of keratinisation, epithelial differentiation, extracellular matrix interaction, and cell adhesion pathways."),
        ("Figure 14", "Hallmark GSEA showing positive enrichment of E2F Targets and G2M Checkpoint and negative enrichment of Epithelial–Mesenchymal Transition and KRAS Signalling in HPV-positive HNSC."),
        ("Figure 15", "Forest plot of hazard ratios and 95% confidence intervals for significant prognostic genes identified via Cox proportional hazards analysis. Although statistically significant, all hazard ratios were close to unity, indicating relatively modest effect sizes."),
        ("Figure 16", "Kaplan-Meier overall survival curves stratified by ZFR2 expression. High-expression group (pink) versus low-expression group (teal); log-rank P = 0.024. Risk table below shows number at risk at each time point.")
    ]
    
    for fig_lbl, fig_desc in fig_legends:
        p_fig = doc.add_paragraph()
        p_fig.paragraph_format.line_spacing = 1.15
        p_fig.paragraph_format.space_after = Pt(8)
        r_flbl = p_fig.add_run(f"{fig_lbl}. ")
        r_flbl.bold = True
        p_fig.add_run(fig_desc)

    # SUPPLEMENTARY FIGURES SECTION
    doc.add_page_break()
    doc.add_heading("Supplementary Figures", level=1)
    
    add_figure(
        doc,
        "HNSC_HPV_GO_Downregulated.png",
        "Supplementary Figure S1. Gene Ontology (GO) biological process enrichment dot plot for downregulated genes in HPV-positive versus HPV-negative HNSC. Terms are ranked by gene ratio, with bubble size indicating gene count and color showing the FDR-adjusted P-value."
    )
    
    add_figure(
        doc,
        "HNSC_HPV_GSEA_ridgeplot.png",
        "Supplementary Figure S2. GSEA ridgeplot of enriched Hallmark gene sets in HPV-positive versus HPV-negative HNSC. The plot displays the distribution of fold changes for genes within each pathway, ranked by normalized enrichment score (NES) and colored by FDR-adjusted P-value."
    )

    # Save document
    output_filename = "Immune_Landscape_HNSC.docx"
    try:
        doc.save(output_filename)
        print(f"Document created successfully: {os.path.abspath(output_filename)}")
    except PermissionError:
        alt_filename = "Immune_Landscape_HNSC_Updated.docx"
        doc.save(alt_filename)
        print(f"WARNING: '{output_filename}' is locked.")
        print(f"Alternative document created successfully: {os.path.abspath(alt_filename)}")

if __name__ == "__main__":
    main()
